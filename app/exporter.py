from docx import Document

def export_answers(results):

    doc = Document()

    doc.add_heading("Questionnaire Answers", level=1)

    for r in results:

        doc.add_heading(r["question"], level=2)

        doc.add_paragraph("Answer: " + r["answer"])

        doc.add_paragraph("Citation: " + r["citation"])

        doc.add_paragraph("Confidence: " + r["confidence"])

    doc.save("answers.docx")